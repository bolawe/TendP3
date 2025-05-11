from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openai import OpenAI
import os
import json

# Initialize client with error handling
try:
    client = OpenAI(api_key=os.environ.get('OPENAI_API_KEY'))
except Exception as e:
    print(f"Failed to initialize OpenAI client: {str(e)}")
    exit(1)

def parse_ai_response(text):
    """Convert AI response to structured data with error handling"""
    try:
        return json.loads(text)  # For JSON responses
    except json.JSONDecodeError:
        # Fallback parsing for text responses
        sections = {}
        current_section = None
        for line in text.split('\n'):
            if ':' in line:
                key, value = line.split(':', 1)
                sections[key.strip()] = value.strip()
        return sections

def generate_report(text):
    """Generate report content with robust error handling"""
    try:
        response = client.chat.completions.create(
            model="gpt-4-turbo",
            response_format={ "type": "json_object" },
            messages=[
                {
                    "role": "system",
                    "content": """You are a tender analysis expert. Return JSON with:
                    {
                        "project_title": str,
                        "key_requirements": [str],
                        "technical_specs": {"requirement": "solution"},
                        "methodology": [str]
                    }"""
                },
                {
                    "role": "user", 
                    "content": text[:20000]  # Safe token limit
                }
            ],
            temperature=0.3
        )
        return parse_ai_response(response.choices[0].message.content)
    except Exception as e:
        print(f"AI generation failed: {str(e)}")
        return None

def create_word_doc(analysis, filename):
    """Create formatted Word document"""
    doc = Document()
    
    # Cover Page
    title = doc.add_paragraph()
    title_run = title.add_run("TECHNICAL PROPOSAL")
    title_run.font.size = Pt(24)
    title_run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Project: {analysis.get('project_title', '')}")
    doc.add_page_break()
    
    # Methodology Section
    doc.add_heading('Technical Methodology', 1)
    for item in analysis.get('methodology', []):
        doc.add_paragraph(item, style='ListBullet')
    
    # Compliance Table
    doc.add_heading('Compliance Matrix', 1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'LightShading-Accent1'
    hdr = table.rows[0].cells
    hdr[0].text = "Requirement"
    hdr[1].text = "Solution"
    hdr[2].text = "Standard"
    
    for req, sol in analysis.get('technical_specs', {}).items():
        row = table.add_row().cells
        row[0].text = req
        row[1].text = sol
        row[2].text = "AWWA C150" if "pipe" in req.lower() else "ISO 9001"
    
    # Save document
    output_path = f"outputs/{filename}_report.docx"
    doc.save(output_path)
    print(f"Report saved to {output_path}")

def main():
    os.makedirs("outputs", exist_ok=True)
    
    for file in os.listdir("outputs"):
        if file.endswith("_cleaned.txt"):
            with open(f"outputs/{file}", "r") as f:
                text = f.read()
            
            analysis = generate_report(text)
            if analysis:
                create_word_doc(analysis, os.path.splitext(file)[0])

if __name__ == "__main__":
    main()
